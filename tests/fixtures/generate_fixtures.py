#!/usr/bin/env python3
"""Generate the RAG ingestion fixture set under tests/fixtures/.

Run from the repository root:
    python tests/fixtures/generate_fixtures.py

The script recreates fixture files in place. It requires reportlab, python-docx,
and pillow for real PDF/DOCX/image fixtures.
"""
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import json
import shutil

ROOT = Path(__file__).resolve().parent


def reset_category_dirs():
    for sub in ['supported_text', 'supported_pdf', 'unsupported', 'odd_names', 'encoding']:
        target = ROOT / sub
        if target.exists():
            shutil.rmtree(target)
        target.mkdir(parents=True, exist_ok=True)


def write_text(path, text, encoding='utf-8', newline=None):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding=encoding, newline=newline)


def write_bytes(path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)


def pdf_with_pages(path, pages):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch

    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    for page_text in pages:
        if page_text is not None:
            text = c.beginText(1 * inch, height - 1 * inch)
            text.setFont('Helvetica', 11)
            for line in page_text.splitlines() or ['']:
                text.textLine(line)
            c.drawText(text)
        c.showPage()
    c.save()


def main():
    reset_category_dirs()

    text_dir = ROOT / 'supported_text'
    write_text(text_dir / 'sample.txt', """RAG ingestion sample text fixture.

This file has multiple lines and normal punctuation.
It is intended for basic plain-text extraction tests.
Source marker: sample.txt
""")
    write_text(text_dir / 'sample.md', """# RAG Fixture Markdown

This markdown file contains headings, bullets, and paragraphs.

## Key points

- Extract headings as normal text.
- Preserve useful list content.
- Ignore no front matter because there is none.

A final paragraph gives the extractor enough content for assertions.
Source marker: sample.md
""")
    write_text(text_dir / 'blank.txt', '')
    write_text(text_dir / 'blank.md', '')
    write_text(text_dir / 'whitespace_only.txt', '   \n\t  \n      ')
    write_text(text_dir / 'whitespace_only.md', '\n\n   \t\n')
    write_text(text_dir / 'leading_trailing_spaces.txt', '   Leading spaces should be stripped.\nMiddle line remains.\nTrailing spaces should be stripped.   \n')
    write_text(text_dir / 'leading_trailing_spaces.md', '   # Trimmed Heading   \n\n   Paragraph with leading and trailing spaces.   \n')
    write_text(text_dir / 'uppercase_extension.TXT', 'Uppercase .TXT extension should be accepted when matching case-insensitively.\n')
    write_text(text_dir / 'uppercase_extension.MD', '# Uppercase Extension\n\nUppercase .MD extension should be accepted.\n')
    write_text(text_dir / 'unicode_text.txt', 'Unicode text: café, naïve, résumé, São Paulo, 東京, مرحبا, emoji: 😀.\n')
    write_text(text_dir / 'special_chars.txt', 'Special characters: ~!@#$%^&*()_+-={}[]|\\:;"\'<>,.?/` and smart quotes “like this”.\n')
    write_text(text_dir / 'single_line.txt', 'A single line fixture for simplest extraction assertions.')
    write_text(text_dir / 'longer_text.txt', """Longer text fixture for chunking-style ingestion tests.

Section 1: Overview
A retrieval-augmented generation system often reads source files, extracts text, chunks it, embeds it, and stores vectors.

Section 2: Details
This fixture is intentionally longer than the other samples but still small enough to inspect quickly in a code review.
It contains repeated terms such as ingestion, extraction, metadata, page number, source name, and validation.

Section 3: Assertions
Tests may assert that the first section, middle section, and final section are all present after ingestion.
The final sentence is unique: fixture sentinel alpaca-river-42.
""")

    pdf_dir = ROOT / 'supported_pdf'
    pdf_with_pages(pdf_dir / 'sample.pdf', ["""PDF sample fixture.
This text appears on page 1.
Source marker: sample.pdf page 1.
Use this for basic PDF text extraction assertions."""])
    pdf_with_pages(pdf_dir / 'multi_page.pdf', [
        """Multi-page PDF fixture.
This text appears on page 1.
Source marker: multi_page.pdf page 1.""",
        """This text appears on page 2.
It supports page-number metadata assertions.
Source marker: multi_page.pdf page 2.""",
        """This text appears on page 3.
Final page sentinel: blue-lantern-17.""",
    ])
    pdf_with_pages(pdf_dir / 'blank.pdf', [None])
    pdf_with_pages(pdf_dir / 'blank_first_page_then_text.pdf', [None, """The first page of this PDF is blank.
This text appears on page 2 only.
Use this to test blank-page skipping and page-number preservation."""])
    pdf_with_pages(pdf_dir / 'text_then_blank_page.pdf', ["""This PDF has text on page 1.
Page 2 is intentionally blank.
Use this to test trailing blank-page handling.""", None])
    pdf_with_pages(pdf_dir / 'whitespace_like_text.pdf', ["     \n\t     \n        "])
    pdf_with_pages(pdf_dir / 'uppercase_extension.PDF', ["""Uppercase .PDF extension fixture.
This should be accepted when extension matching is case-insensitive."""])

    unsup = ROOT / 'unsupported'
    from docx import Document
    doc = Document()
    doc.add_heading('Unsupported DOCX fixture', level=1)
    doc.add_paragraph('This file is intentionally unsupported by the current ingestion app.')
    doc.save(str(unsup / 'sample.docx'))
    write_text(unsup / 'sample.csv', 'id,name,note\n1,Ada,CSV should be rejected as unsupported\n')
    write_text(unsup / 'sample.json', json.dumps({'purpose': 'unsupported format fixture', 'should_ingest': False}, indent=2) + '\n')
    write_text(unsup / 'sample.html', '<!doctype html><html><body><h1>Unsupported HTML</h1><p>Should be rejected.</p></body></html>\n')
    write_text(unsup / 'sample.py', 'print("Unsupported Python source fixture")\n')
    from PIL import Image
    img = Image.new('RGB', (1, 1), color=(255, 255, 255))
    img.save(unsup / 'sample.png')
    img.save(unsup / 'sample.jpg', 'JPEG')
    with ZipFile(unsup / 'sample.zip', 'w', ZIP_DEFLATED) as z:
        z.writestr('inside.txt', 'This archive should be rejected as unsupported.\n')
    write_bytes(unsup / 'sample.exe', b'MZ\x90\x00placeholder unsupported exe bytes\x00')
    write_bytes(unsup / 'sample.bin', b'\x00\x01\x02\xff\xfeunsupported binary fixture\x00')

    odd = ROOT / 'odd_names'
    write_text(odd / 'file with spaces.txt', 'Filename contains spaces. Source marker: file with spaces.txt\n')
    write_text(odd / 'mixed.Case.Md', '# Mixed Case Markdown\n\nExtension has mixed case and should be handled case-insensitively.\n')
    write_text(odd / 'notes.v1.txt', 'Filename contains an extra dot before the extension. Version marker: v1.\n')
    write_text(odd / 'README.TEST.MD', '# README Test\n\nFilename contains multiple uppercase segments.\n')
    write_text(odd / 'weird-name_123.txt', 'Filename contains hyphen, underscore, and digits.\n')

    enc = ROOT / 'encoding'
    write_text(enc / 'utf8_accents.txt', 'Accented UTF-8: crème brûlée, jalapeño, façade, coöperate, voilà.\n')
    write_text(enc / 'symbols_punctuation.txt', 'Symbols and punctuation: © ® ™ § ¶ • – — … ± × ÷ ≈ ≠ ≤ ≥ ∑ ∞.\n')
    write_text(enc / 'newlines_only.txt', '\n\n\n\n')
    write_text(enc / 'tabs_and_spaces.txt', '\tIndented with a tab\n    Indented with four spaces\n\t    Mixed tab and spaces\n')
    write_bytes(enc / 'invalid_utf8_bytes.txt', b'Valid prefix. Invalid bytes: \xff\xfe\xfa. Valid suffix.\n')

    print(f'Generated fixtures in {ROOT}')


if __name__ == '__main__':
    main()
